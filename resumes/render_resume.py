#!/usr/bin/env python3
"""
Render a structured resume JSON into a clean, professional .docx.

Usage:
    python3 render_resume.py master.json
    python3 render_resume.py variants/vc-investor.json -o output/AdamBreaux_VC.docx

Design goals: Calibri family, 1" margins, large bold name, ruled section
headers, bold lead-ins on bullets. Approximates the look of the original
resume while staying consistent. Edit the data files, never hand-edit the docx.
"""
import argparse
import json
import os
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches, RGBColor

ACCENT = RGBColor(0x1F, 0x2A, 0x44)   # deep navy for headers
GRAY = RGBColor(0x44, 0x44, 0x44)
FONT = "Calibri"


def _set_margins(doc):
    for s in doc.sections:
        s.top_margin = Inches(0.6)
        s.bottom_margin = Inches(0.6)
        s.left_margin = Inches(0.8)
        s.right_margin = Inches(0.8)


def _no_space(p, before=0, after=2, line=1.0):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def _bottom_border(p):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F2A44")
    pbdr.append(bottom)
    pPr.append(pbdr)


def _run(p, text, size=10.5, bold=False, italic=False, color=None):
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    if color is not None:
        r.font.color.rgb = color
    return r


def render(data, out_path):
    doc = Document()
    _set_margins(doc)
    doc.styles["Normal"].font.name = FONT
    doc.styles["Normal"].font.size = Pt(10.5)

    # Name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _no_space(p, after=1)
    _run(p, data["name"], size=24, bold=True, color=ACCENT)

    # Contact
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _no_space(p, after=6)
    _run(p, data.get("contact", ""), size=9.5, color=GRAY)

    # Summary
    if data.get("summary"):
        p = doc.add_paragraph()
        _no_space(p, after=6)
        _run(p, data["summary"], size=10.5)

    for section in data.get("sections", []):
        # Section header
        p = doc.add_paragraph()
        _no_space(p, before=6, after=3)
        _bottom_border(p)
        _run(p, section["heading"].upper(), size=11.5, bold=True, color=ACCENT)

        stype = section.get("type")
        if stype == "skills":
            p = doc.add_paragraph()
            _no_space(p, after=4)
            items = section["items"]
            for i, it in enumerate(items):
                _run(p, it, size=10.5)
                if i < len(items) - 1:
                    _run(p, "   •   ", size=10.5, color=ACCENT)
            continue

        for item in section["items"]:
            # Title line
            p = doc.add_paragraph()
            _no_space(p, before=4, after=0)
            _run(p, item["title"], size=11, bold=True)
            if item.get("dates"):
                tab = p.add_run("\t")
                pf = p.paragraph_format
                # right-align tab stop
                from docx.enum.text import WD_TAB_ALIGNMENT
                from docx.shared import Inches as In
                pf.tab_stops.add_tab_stop(In(6.9), WD_TAB_ALIGNMENT.RIGHT)
                _run(p, item["dates"], size=10, italic=True, color=GRAY)
            # Org line
            org_bits = []
            if item.get("org"):
                org_bits.append(item["org"])
            if item.get("location"):
                org_bits.append(item["location"])
            if org_bits:
                p = doc.add_paragraph()
                _no_space(p, after=2)
                _run(p, "  |  ".join(org_bits), size=10, italic=True, color=GRAY)
            # Bullets
            for b in item.get("bullets", []):
                p = doc.add_paragraph(style="List Bullet")
                _no_space(p, after=2, line=1.0)
                p.paragraph_format.left_indent = Inches(0.25)
                # bold lead-in if "Lead-in: rest" pattern present
                if ": " in b and len(b.split(": ")[0]) < 40:
                    head, rest = b.split(": ", 1)
                    _run(p, head + ": ", size=10.5, bold=True)
                    _run(p, rest, size=10.5)
                else:
                    _run(p, b, size=10.5)

    os.makedirs(os.path.dirname(out_path) or ".", exist_ok=True)
    doc.save(out_path)
    return out_path


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("data", help="path to resume JSON")
    ap.add_argument("-o", "--out", help="output .docx path")
    args = ap.parse_args()

    with open(args.data) as f:
        data = json.load(f)

    if args.out:
        out = args.out
    else:
        base = os.path.splitext(os.path.basename(args.data))[0]
        out = os.path.join("output", f"AdamBreaux_{base}.docx")

    path = render(data, out)
    print(f"Wrote {path}")


if __name__ == "__main__":
    main()
